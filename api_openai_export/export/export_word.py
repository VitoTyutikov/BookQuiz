import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_question_doc(json_input) -> str:
    # Load JSON data
    
    # book_json = json.dumps(json_input)
    # data = json.loads(json_input)
    data = json_input

    # Create a new Document
    doc = Document()
    doc.add_heading(data['bookTitle'],
                    level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sort chapters by chapterId
    chapters_sorted = sorted(data['chapters'], key=lambda x: x['chapterId'])

    # Question numbering across all chapters
    # question_number = 1

    # Iterate over each chapter
    for chapter in chapters_sorted:
        doc.add_heading(
            f"{chapter['chapterTitle']}", level=1)

        # Iterate over each question
        for question in chapter['questions']:
            doc.add_paragraph(
                f"{question['questionText']}", style='ListNumber')
            # question_number += 1

            # List answers with indentation and numbering
            for idx, answer in enumerate(question['answers'], start=1):
                para = doc.add_paragraph(f"{idx}. {answer['answerText']}")
                para.paragraph_format.left_indent = Pt(36)

    # Add a section for correct answers
    doc.add_page_break()
    doc.add_heading('List of Correct Answers', level=1)

    # Reset question numbering for the answer list
    question_number = 1
    for chapter in chapters_sorted:
        doc.add_heading(chapter['chapterTitle'], level=2)
        for question in chapter['questions']:
            correct_answer_idx = [idx for idx, answer in enumerate(
                question['answers'], start=1) if answer['isCorrect']][0]
            doc.add_paragraph(
                f"Question {question_number} - Correct Answer: {correct_answer_idx}")
            question_number += 1

    # Save the document
    directory = "./generated_word_files"
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Sanitize filename to remove disallowed characters and limit length
    sanitized_title = "".join(char for char in data['bookTitle'] if char.isalnum() or char in [" ", "_"]).rstrip()
    file_name = f"docx_questions_{data['bookId']}_{sanitized_title}.docx"
    
    # Save the document
    full_path = os.path.join(directory, file_name)
    doc.save(full_path)
    return full_path


